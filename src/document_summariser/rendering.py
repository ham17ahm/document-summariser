from __future__ import annotations

from html import escape
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from document_summariser.config import AppConfig


class DocxRenderer:
    def render(self, summary: str, output_path: Path, config: AppConfig) -> None:
        try:
            self._render_with_python_docx(summary, output_path, config)
        except ModuleNotFoundError:
            self._render_minimal_docx(summary, output_path, config)

    def _render_with_python_docx(self, summary: str, output_path: Path, config: AppConfig) -> None:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        document = Document()
        title_text = str(config.output.get("title", "Urdu Summary"))
        font = str(config.output.get("font", "Noto Nastaliq Urdu"))
        style = document.styles["Normal"]
        style.font.name = font
        _set_style_complex_script_font(style, font)

        title = document.add_heading(title_text, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_paragraph_rtl(title)
        for run in title.runs:
            _set_run_rtl(run, font)

        for block in summary.split("\n\n"):
            if not block.strip():
                continue
            paragraph = document.add_paragraph(block)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_paragraph_rtl(paragraph)
            for run in paragraph.runs:
                _set_run_rtl(run, font)

        document.save(output_path)

    def _render_minimal_docx(self, summary: str, output_path: Path, config: AppConfig) -> None:
        font = str(config.output.get("font", "Noto Nastaliq Urdu"))
        title_text = str(config.output.get("title", "Urdu Summary"))
        paragraphs = [title_text, *summary.split("\n\n")]
        body = "".join(_docx_paragraph(paragraph, font) for paragraph in paragraphs if paragraph.strip())
        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f"<w:body>{body}<w:sectPr/></w:body></w:document>"
        )
        with ZipFile(output_path, "w", ZIP_DEFLATED) as archive:
            archive.writestr(
                "[Content_Types].xml",
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
                '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
                '<Default Extension="xml" ContentType="application/xml"/>'
                '<Override PartName="/word/document.xml" '
                'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
                "</Types>",
            )
            archive.writestr(
                "_rels/.rels",
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                '<Relationship Id="rId1" '
                'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
                'Target="word/document.xml"/>'
                "</Relationships>",
            )
            archive.writestr("word/document.xml", document_xml)


def _docx_paragraph(text: str, font: str) -> str:
    escaped = escape(text, quote=True)
    escaped_font = escape(font, quote=True)
    return (
        "<w:p><w:pPr><w:bidi/><w:jc w:val=\"right\"/></w:pPr>"
        "<w:r><w:rPr>"
        f"<w:rFonts w:ascii=\"{escaped_font}\" w:hAnsi=\"{escaped_font}\" w:cs=\"{escaped_font}\"/>"
        "<w:rtl/>"
        f"</w:rPr><w:t>{escaped}</w:t></w:r></w:p>"
    )


def _set_paragraph_rtl(paragraph) -> None:
    from docx.oxml import OxmlElement

    paragraph_format = getattr(paragraph, "paragraph_format", None)
    if paragraph_format is not None and hasattr(paragraph_format, "right_to_left"):
        paragraph_format.right_to_left = True

    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find("w:bidi", p_pr.nsmap) is None:
        p_pr.append(OxmlElement("w:bidi"))


def _set_run_rtl(run, font: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run.font.name = font
    r_pr = run._r.get_or_add_rPr()

    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attribute in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(attribute), font)

    if r_pr.find("w:rtl", r_pr.nsmap) is None:
        r_pr.append(OxmlElement("w:rtl"))


def _set_style_complex_script_font(style, font: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attribute in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(attribute), font)
